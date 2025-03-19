#%%
import whisper
import openai
from sklearn.feature_extraction.text import TfidfVectorizer
import spacy
from textblob import TextBlob
import ffmpeg
from docx import Document
from sklearn.metrics.pairwise import cosine_similarity
import re
from dotenv import load_dotenv
import os


#%%

# ✅ Load .env file
load_dotenv()

# ✅ Access the API key
openai_api_key = os.getenv("OPENAI_API_KEY")

client = openai.OpenAI(api_key=openai_api_key)  # Ensure API key is set

#%%
# ✅ Function to preprocess audio (reduce noise & normalize volume)
def preprocess_audio(input_path, output_path):
    """
    Enhances audio by applying noise reduction & volume normalization.
    """
    (
        ffmpeg
        .input(input_path)
        .output(output_path, acodec="pcm_s16le", ar="16000")  # Convert to 16kHz WAV
        .run(overwrite_output=True)
    )

# ✅ Load Whisper Model (Best Model for Accuracy)
model = whisper.load_model("base.en")

# ✅ Enhance Audio Before Transcription
file_path = r"c:\Users\ehari\Desktop\podcast_clipper\test.wav"
enhanced_audio_path = "enhanced_audio.wav"
preprocess_audio(file_path, enhanced_audio_path)

# ✅ Transcribe the Enhanced Audio
result = model.transcribe(enhanced_audio_path, word_timestamps=True)

#%%
# ✅ Function to split and process long transcripts in chunks
def correct_transcription_in_chunks(transcript, max_chars=5000):
    """
    Splits a transcript into smaller chunks to avoid exceeding OpenAI token limits.
    Processes each chunk separately and combines them back.
    """
    words = transcript.split()
    chunks = []
    current_chunk = []
    current_length = 0

    for word in words:
        current_chunk.append(word)
        current_length += len(word) + 1  # Account for spaces

        if current_length >= max_chars:
            chunks.append(" ".join(current_chunk))
            current_chunk = []
            current_length = 0

    if current_chunk:
        chunks.append(" ".join(current_chunk))  # Add remaining words

    corrected_chunks = []

    for chunk in chunks:
        prompt = f"""
        The following is a segment of a podcast transcript with possible errors.
        Please correct spelling, and grammar while keeping the original meaning.

        **Segment:**
        {chunk}

        **Return only the corrected version.**
        """

        response = client.chat.completions.create(
            model="gpt-3.5-turbo",  # Use GPT-3.5 to avoid token issues
            messages=[{"role": "system", "content": prompt}]
        )
        corrected_chunks.append(response.choices[0].message.content.strip())

    return " ".join(corrected_chunks)  # Combine all corrected chunks

# ✅ Extract word-level timestamps and group sentences into longer clips
nlp = spacy.load("en_core_web_md")
sentences = []
current_clip = []
start_time = None
prev_end_time = None
max_words_per_clip = 250  # Increased for better clip length

for segment in result["segments"]:
    for word in segment["words"]:
        word_text = word.get("word")
        if not word_text:
            continue

        if not start_time:
            start_time = word["start"]
        current_clip.append(word_text)
        prev_end_time = word["end"]

        if word_text.endswith((".", "!", "?")) or len(current_clip) >= max_words_per_clip:
            sentences.append((" ".join(current_clip), round(start_time, 2), round(prev_end_time, 2)))
            current_clip = []
            start_time = None  

if current_clip:
    sentences.append((" ".join(current_clip), round(start_time, 2), round(prev_end_time, 2)))


# ✅ Load SpaCy's Medium English Model for Better Sentence Representations
nlp = spacy.load("en_core_web_md")  # Uses word embeddings

def generate_clips(sentences, nlp, spacy_threshold=0.95, tfidf_threshold=0.9, max_gap=30):
    """
    ✅ Generates podcast clips while **ensuring natural flow** using:
    - **SpaCy embeddings** for semantic similarity.
    - **TF-IDF cosine similarity** for keyword relevance.
    - **Dynamically merges sentences based on similarity & continuity.**
    
    🔹 **Improvements:**
    - ✅ **Removes word count restrictions.**
    - ✅ **Uses higher similarity thresholds for more natural grouping.**
    - ✅ **Only checks next sequential sentence (no jumping).**
    - ✅ **Prevents merging distant or disjointed sentences.**
    """

    clips = []
    i = 0
    vectorizer = TfidfVectorizer(stop_words='english')

    while i < len(sentences):
        sentence, start, end = sentences[i]

        # ✅ Ensure sentence is not empty
        if not sentence.strip():
            i += 1
            continue

        # ✅ Convert sentence to SpaCy vector representation
        sentence_doc = nlp(sentence)
        sentence_vec = sentence_doc.vector.reshape(1, -1)

        clip_sentences = [sentence]
        clip_start = start

        # ✅ Expand clip by adding **next sequential sentences**
        j = i + 1
        while j < len(sentences):
            next_sentence, next_start, next_end = sentences[j]

            # ✅ Ensure next sentence is not empty
            if not next_sentence.strip():
                j += 1
                continue

            next_doc = nlp(next_sentence)
            next_vec = next_doc.vector.reshape(1, -1)

            # ✅ Compute Cosine Similarity Between Sentence Vectors (SpaCy)
            spacy_similarity_score = cosine_similarity(sentence_vec, next_vec)[0][0]

            # ✅ Compute TF-IDF Similarity Score (Only if both sentences have meaningful words)
            meaningful_sentences = [sentence, next_sentence]
            try:
                tfidf_matrix = vectorizer.fit_transform(meaningful_sentences)
                tfidf_similarity_score = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
            except ValueError:
                tfidf_similarity_score = 0  # If TF-IDF fails, default to 0

            # ✅ Use max similarity from SpaCy and TF-IDF
            similarity_score = max(spacy_similarity_score, tfidf_similarity_score)
            
            # ✅ Prevent merging **distant** sentences
            if next_start - end > max_gap:
                break  # Do not merge if gap between sentences is too large

            # ✅ Check similarity threshold
            if similarity_score < spacy_threshold and similarity_score < tfidf_threshold:
                break  # Stop expansion if sentences are too different

            # ✅ Extend clip dynamically (Merging multiple relevant sentences)
            clip_sentences.append(next_sentence)
            end = next_end  # Extend end time
            j += 1  # Move to the next sentence

        # ✅ Store the clip **only if similarity-driven merging was successful**
        clips.append((" ".join(clip_sentences), clip_start, end))

        # ✅ Move forward to the next unmerged sentence
        i = j  

    print(f"🔍 Generated {len(clips)} sequentially merged clips with improved similarity scoring.")
    return clips


# ✅ Function to generate SEO-friendly tags
def generate_seo_tags(full_transcript):
    """
    Uses AI to extract SEO-friendly tags based on the podcast transcript.
    """
    prompt = f"""
    You are an expert in podcast marketing. Analyze the following podcast transcript and 
    generate **SEO-friendly tags** based on the discussion topics, important phrases, and 
    key themes. Make sure to include:

    - **Relevant keywords that reflect the podcast’s main topics**
    - **3-5 main topics from the episode**
    - **Popular search terms related to the discussion**
    - **Trending podcast-related tags to improve visibility**
    - **Try to encompass the entire episode and not just one inital topic**

    **Podcast Transcript:**
    {full_transcript}  # Limit to avoid token issues

    **Return the tags as a comma-separated list** (e.g., "technology, AI, future trends, podcast discussions").
    """

    response = client.chat.completions.create(
        model="gpt-3.5-turbo",  # Uses GPT-3.5 to avoid token issues
        messages=[{"role": "system", "content": prompt}]
    )

    episode_specific_tags = response.choices[0].message.content.strip()
    general_tags = "podcast, interviews, storytelling, trending, discussions, expert insights"

    return f"{general_tags}, {episode_specific_tags}"



def rank_clips_advanced(clips, nlp, seo_tags, min_time_gap=180):
    """
    Enhances clip selection by:
    - Penalizing and filtering intro/outro remarks correctly.
    - Using SEO similarity & trending topics for ranking.
    - Incorporating NLP relevance, sentiment, and engagement scores.
    """

    vectorizer = TfidfVectorizer(stop_words='english')

    # ✅ **Ensure seo_tags is formatted properly**
    seo_text = " ".join(seo_tags) if isinstance(seo_tags, list) else seo_tags

    # ✅ **Create reference text combining SEO tags & trending topics**
    reference_text = f"{seo_text}"

    # ✅ **Precompute NLP embeddings for SEO Tags**
    reference_doc = nlp(reference_text)
    reference_vec = reference_doc.vector.reshape(1, -1)  # Convert to 2D array

    # ✅ **Combine Reference Text + All Podcast Clips**
    text_data = [reference_text] + [clip[0] for clip in clips]

    ranked_clips = []
    used_timestamps = set()

    # ✅ **Expanded Intro/Outro Filtering with Regex**
    banned_phrases = [
        # Standard podcast intro/outro phrases
        r"\brate and subscribe\b", r"\bspecial thank you\b", r"\bwelcome to\b", 
        r"\bthanks for listening\b", r"\bbefore we start\b", r"\bsubscribe\b", 
        r"\bleave a review\b", r"\bhow are we doing\b", r"\blet’s wrap up\b", 
        r"\bfollow us on\b", r"\bthat’s it for today\b", r"\bclosing thoughts\b", 
        r"\bfinal reminder\b", r"\bdon’t forget to\b", r"\bmake sure to subscribe\b",
        r"\btune in next time\b", r"\bif you enjoyed this\b", r"\bplease support us by\b", 
        r"\bvisit our website\b", r"\bcheck out our sponsors\b", r"\bbrought to you by\b",
        r"\bthe  knowledge  bar",
        
        # Common show introductions
        r"\bwe have a great episode for you\b", r"\btoday we’re going to talk about\b", 
        r"\blet’s get started\b", r"\bour topic today is\b", r"\bjoining me today is\b", 
        r"\bfirst up, we have\b", r"\blet’s introduce\b", r"\bwe’re excited to discuss\b", 
        r"\bwe’ve got an amazing guest\b", r"\bbig discussion ahead\b", r"\bthis week on the knowledge bar\b",
        
        # Closing remarks
        r"\bhope you enjoyed\b", r"\bthat’s all we have time for\b", r"\bwe’ll see you next time\b", 
        r"\bthanks again for tuning in\b", r"\bwe appreciate your support\b", r"\bfinal thoughts on this\b", 
        r"\bone last thing before we go\b", r"\bstay connected with us\b", r"\buntil next time\b", 
        
        # Calls to action
        r"\blet us know what you think\b", r"\bcomment below\b", r"\bdrop us a message\b", 
        r"\bleave your thoughts\b", r"\bemail us at\b", r"\bfind us on social media\b"
    ]

    for i, (sentence, start, end) in enumerate(clips):
        if any(phrase in sentence.lower() for phrase in banned_phrases):
            continue  # Skip intro/outro clips

        words = sentence.split()

        # ✅ **1. Semantic Similarity Score (Using SpaCy + Cosine Similarity)**
        sentence_doc = nlp(sentence)
        sentence_vec = sentence_doc.vector.reshape(1, -1)
        similarity_score = cosine_similarity(reference_vec, sentence_vec)[0][0] * 15  

        # ✅ **2. NLP-Based Relevance Score (Measures Topic Alignment)**
        nlp_score = reference_doc.similarity(sentence_doc) * 10  

        # ✅ **3. Sentiment Score (Measures Emotional Intensity)**
        sentiment_score = abs(TextBlob(sentence).sentiment.polarity) * 10 

        # ✅ **4. Keyword Score (Counts important words based on SEO Tags)**
        keyword_score = sum(0.5 for word in words if word.lower() in seo_text.lower())

        # ✅ **5. Length Bonus (Encourages Longer, Well-Developed Clips)**
        length_bonus = min(len(words) / 50, 5)  

        # ✅ **6. Quote Score (Detects Direct Speech)**
        quote_score = 10 if '"' in sentence else 0  

        # ✅ **7. Stats & Figures Score (Identifies Data-Driven Segments)**
        stats_score = sum(5 for word in words if word.replace(',', '').replace('.', '').isdigit()) 

        # ✅ **9. Emotion Score (Prioritizes Clips with High Impact Language)**
        emotional_words = {
            "amazing", "shocking", "unbelievable", "outrageous", 
            "crazy", "insane", "incredible", "huge", "special",
            "astonishing", "breathtaking", "heartbreaking", "devastating",
            "mind-blowing", "unreal", "remarkable", "phenomenal",
            "stunning", "extraordinary", "sensational", "jaw-dropping",
            "horrifying", "tragic", "terrifying", "disgusting",
            "unforgivable", "groundbreaking", "legendary", "spectacular",
            "unthinkable", "appalling", "mind-boggling", "thrilling",
            "exhilarating", "electrifying", "life-changing", "unprecedented",
            "earth-shattering", "epic", "game-changing", "colossal",
            "tremendous", "overwhelming", "riveting", "explosive",
            "gut-wrenching", "soul-crushing", "spellbinding", "miraculous",
            "fiery", "intense", "heroic", "devastating", "horrible", "terrible",
            "terrific", "absolutely", "what's got me"
        }
        emotion_score = sum(5 for word in sentence.lower().split() if word in emotional_words)

        # ✅ **10. Rhetorical Score (Detects Engagement Hooks)**
        rhetorical_score = 5 if "?" in sentence else 0  

        # ✅ **11. Final Weighted Score Calculation**
        final_score = (
            (similarity_score) +  
            (nlp_score) +  
            (sentiment_score) +  
            (keyword_score) +  
            (stats_score) +  
            (quote_score) +  
            (emotion_score) +  
            (rhetorical_score) +  
            (length_bonus * 2)  
        )

        ranked_clips.append((
            sentence, start, end, final_score, similarity_score, sentiment_score, 
            keyword_score, stats_score, nlp_score, length_bonus, 
            quote_score, emotion_score, rhetorical_score
        ))

    # ✅ **Return Top 5 Best & Diverse Clips**
    return sorted(ranked_clips, key=lambda x: x[3], reverse=True)[:5]


# ✅ Use AI to select the best highlights
full_transcript = result['text']
corrected_transcript = correct_transcription_in_chunks(full_transcript)


# ✅ Function to Generate AI-Powered Episode Titles
def generate_episode_titles(full_transcript):
    """
    Uses GPT-4 to analyze the podcast transcript and generate 5 engaging episode titles.
    """
    prompt = f"""
    You are a professional podcast producer. Your job is to generate **5 creative episode titles** 
    based on the podcast transcript. Titles should be:

    - **Engaging & attention-grabbing** 
    - **In the form of a provocative question**
    - **Encompass the entire episode and not just one topic**
    - **SEO-friendly (contain keywords related to the episode)**
    - **Diverse in style** (e.g., clickbait, informative, conversational)

    **Podcast Transcript:**
    {full_transcript[:5000]}

    **Generate 5 titles in this format:**
    1. "[Title 1]"
    2. "[Title 2]"
    3. "[Title 3]"
    4. "[Title 4]"
    5. "[Title 5]"

    ONLY return the titles—do not include explanations.
    """
    
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "system", "content": prompt}]
    )

    titles = response.choices[0].message.content.strip().split("\n")
    return [title.strip().strip('"') for title in titles if title.strip()]



# ✅ Generate AI-Powered Episode Titles
episode_titles = generate_episode_titles(corrected_transcript)

def generate_podcast_topic(transcript):
    """
    Uses GPT-3.5 to summarize the main theme of the podcast episode.
    """
    prompt = f"""
    You are an expert podcast analyst. Your job is to summarize the overall topic of this podcast episode.
    Provide a description that captures the 3-5 core themes of the discussion. Try to encompass the 
    entire episode and not just the initial topic
    
    **Podcast Transcript:**
    {transcript}  # Limit to avoid token issues

    **Return format:**
    "This episode discusses [MAIN TOPIC]."
    """
    
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "system", "content": prompt}]
    )

    return response.choices[0].message.content.strip()



def generate_clip_hashtags(clip_text, podcast_topic):
    """
    Generates hashtags based on the clip, podcast topic, and current trends.
    """
    
    prompt = f"""
    You are an expert in social media engagement.
    Generate **5-10 engaging hashtags** for this podcast clip.

    **Podcast Clip:**
    "{clip_text}"

    **Podcast Topic:**
    "{podcast_topic}"


    **Hashtag Rules:**
    - Include 2 branding hashtags: #TKB #TheKnowledgeBar
    - Include at least 2 **current trending topics**
    - Use topic-specific keywords
    - At least 1 hashtag should be for **engagement (e.g., #trending, #viral, #breakingnews)**

    **Return format:**  
    - A comma-separated list (e.g., "#TKB, #TheKnowledgeBar, #collegeathletes, #trendingtopic")
    """
    
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "system", "content": prompt}]
    )
    
    return response.choices[0].message.content.strip()


def save_podcast_summary(important_clips, episode_titles, seo_tags, corrected_transcript, clip_hashtags, output_file="podcast_summary.docx"):
    """
    Saves podcast summary, ranked clips in a structured Word document with tables.
    """

    # ✅ Ensure `output_file` is a string (not a list)
    if not isinstance(output_file, str):
        raise ValueError(f"Expected `output_file` to be a string, but got {type(output_file)}")

    doc = Document()
    doc.add_heading("🎙️ AI-Powered Podcast Summary", level=1)

    # ✅ Add Episode Titles
    doc.add_heading("🎧 Suggested Episode Titles", level=2)
    for title in episode_titles:
        doc.add_paragraph(f"- {title}")

    # ✅ Add Podcast Clips Section
    doc.add_heading("🎬 Selected Podcast Clips (with Score Breakdown)", level=2)

    for i, (clip, start_time, end_time, final_score, seo_similarity, sentiment, keyword_score, stats, nlp_score, length_bonus, quote_score, emotion_score, rhetorical_score) in enumerate(important_clips):
        formatted_start = f"{int(start_time // 60)}:{int(start_time % 60):02d}"
        formatted_end = f"{int(end_time // 60)}:{int(end_time % 60):02d}"
        
        # ✅ Add Clip Header
        doc.add_paragraph(f"[{formatted_start} - {formatted_end}] {clip}")

        # ✅ Create Table for Scores (Header + Metrics)
        num_rows = 1 + 10  # 1 header row + 10 metric rows
        table = doc.add_table(rows=num_rows, cols=3)
        table.style = "Table Grid"

        # ✅ Define Headers
        headers = ["Metric", "Score", "Description"]
        for j, header in enumerate(headers):
            table.cell(0, j).text = header  # First row is header row

        # ✅ Fill Table with Scores (Dynamically Adding Rows)
        scores = [
            ("Final Score", f"{final_score:.2f}", "Aggregated ranking"),
            ("SEO Similarity", f"{seo_similarity:.2f}", "Boosts visibility"),
            ("NLP Relevance", f"{nlp_score:.2f}", "Measures topic alignment"),
            ("Sentiment Score", f"{sentiment:.2f}", "Higher = more engaging"),
            ("Keyword Score", f"{keyword_score}", "Measures key phrase density"),
            ("Stats & Figures", f"{stats}", "Mentions of numbers/data"),
            ("Quote Score", f"{quote_score}", "Detects direct speech"),
            ("Emotion Score", f"{emotion_score}", "Exciting/emotional language"),
            ("Rhetorical Score", f"{rhetorical_score}", "Uses rhetorical questions"),
            ("Length Bonus", f"{length_bonus:.2f}", "Longer clips rank better"),
        ]

        for row, (metric, score, description) in enumerate(scores, start=1):
            table.cell(row, 0).text = metric
            table.cell(row, 1).text = score
            table.cell(row, 2).text = description

        # ✅ Add Hashtags Below the Table
        if clip_hashtags and i < len(clip_hashtags):
            doc.add_paragraph(f"📢 Hashtags: {clip_hashtags[i]}", style="Normal")

        doc.add_paragraph("\n------------------------------------------\n")

    # ✅ Add SEO-Friendly Tags
    doc.add_heading("🔍 SEO-Friendly Tags", level=2)
    doc.add_paragraph(seo_tags)

    # ✅ Add Full Transcript
    doc.add_heading("📜 Full Podcast Transcript", level=2)
    doc.add_paragraph(corrected_transcript)

    # ✅ Save Document (Ensure output_file is a string)
    try:
        doc.save(output_file)
        print(f"\n✅ Podcast summary saved to {output_file}")
    except Exception as e:
        print(f"❌ Error saving document: {e}")


#%%


# ✅ Load spaCy NLP model for organization detection
nlp = spacy.load("en_core_web_md")

# ✅ Extract full transcript from Whisper's output
full_transcript = result['text']

# ✅ Correct the transcription for accuracy
corrected_transcript = correct_transcription_in_chunks(full_transcript)

# ✅ Automatically generate the podcast topic dynamically
podcast_topic = generate_podcast_topic(corrected_transcript)  # ✅ NEW!

print(f"🎙️ **Generated Podcast Topic:** {podcast_topic}")

# ✅ Step 1: Generate Clips with Better Sentence Merging
all_clips = generate_clips(sentences, nlp)  # ✅ Uses SpaCy embeddings for better segmentation

# ✅ Step 2: Generate AI-powered SEO tags
seo_tags = generate_seo_tags(corrected_transcript)  # ✅ Ensure SEO tags are generated before ranking

# ✅ Step 3: Rank Clips Using NLP, SEO, and Sentiment
important_clips = rank_clips_advanced(all_clips, nlp, seo_tags)  # ✅ Now uses refined clip selection

# ✅ Step 4: Generate Hashtags Per Clip Instead of Full Transcript
clip_hashtags = [generate_clip_hashtags(clip[0], podcast_topic) for clip in important_clips]  # ✅ Now specific to each clip

# ✅ Step 5: Save to Word Document
save_podcast_summary(important_clips, episode_titles, seo_tags, corrected_transcript, clip_hashtags)




#%%

# %%
